"""
Module d'exports et rapports SPIC 2.0
Export Word, Excel, PDF avec templates personnalisables
Cohérent avec l'architecture refactorisée
"""

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime, date
from typing import Dict, List, Optional
import io
import base64
import logging

import config
from utils import *
from utils.form_utils import FormComponents
from utils.db_utils import DatabaseUtils
from utils.constants import UI_MESSAGES, CACHE_CONFIG

logger = logging.getLogger(__name__)

def show_exports(db_manager, user_session: Dict):
    """Interface principale d'exports et rapports"""
    
    st.header("📥 Exports et Rapports")
    
    # Navigation par onglets (cohérent avec autres modules)
    tab1, tab2, tab3, tab4 = st.tabs([
        "📊 Exports Standard",
        "📝 Rapports Personnalisés",
        "📋 Templates",
        "📈 Exports Graphiques"
    ])
    
    db_utils = DatabaseUtils(db_manager)
    
    with tab1:
        show_standard_exports(db_manager, db_utils, user_session)
    
    with tab2:
        show_custom_reports(db_manager, db_utils, user_session)
    
    with tab3:
        show_templates_management(db_manager, db_utils, user_session)
    
    with tab4:
        show_graphic_exports(db_manager, db_utils, user_session)

def show_standard_exports(db_manager, db_utils, user_session: Dict):
    """Exports standard (Excel, Word, CSV)"""
    
    st.subheader("📊 Exports Standard")
    
    # Sélection du type d'export
    export_type = st.selectbox(
        "Type d'export",
        options=[
            "portfolio", "operation_detail", "journal", "risks", 
            "budget_evolution", "rem_summary", "alerts"
        ],
        format_func=lambda x: {
            "portfolio": "📋 Portefeuille d'opérations",
            "operation_detail": "🎯 Détail d'une opération",
            "journal": "📖 Journal des modifications",
            "risks": "🚨 Analyse des risques",
            "budget_evolution": "💰 Évolution budgétaire",
            "rem_summary": "📈 Synthèse REM",
            "alerts": "⚠️ Alertes actives"
        }[x],
        key="standard_export_type"
    )
    
    # Configuration selon le type
    export_config = configure_export(export_type, db_utils, user_session)
    
    if not export_config:
        return
    
    # Formats disponibles
    available_formats = get_available_formats(export_type)
    
    col1, col2 = st.columns(2)
    
    with col1:
        selected_format = st.selectbox(
            "Format d'export",
            options=available_formats,
            format_func=lambda x: {
                "excel": "📊 Excel (.xlsx)",
                "word": "📄 Word (.docx)",
                "csv": "📋 CSV (.csv)",
                "pdf": "📑 PDF (.pdf)"
            }.get(x, x),
            key="export_format"
        )
    
    with col2:
        include_charts = st.checkbox(
            "Inclure les graphiques",
            value=True,
            key="include_charts"
        ) if selected_format in ["word", "pdf"] else False
    
    # Prévisualisation des données
    if st.checkbox("👁️ Aperçu des données", key="preview_data"):
        show_export_preview(db_manager, export_type, export_config, user_session)
    
    # Bouton d'export
    if st.button(f"📥 Générer l'export {selected_format.upper()}", type="primary"):
        try:
            with st.spinner("Génération de l'export en cours..."):
                
                # Générer l'export selon le format
                if selected_format == "excel":
                    file_data, filename = generate_excel_export(
                        db_manager, export_type, export_config, user_session
                    )
                    mime_type = config.FORMATS_EXPORT["excel"]["mime"]
                
                elif selected_format == "word":
                    file_data, filename = generate_word_export(
                        db_manager, export_type, export_config, user_session, include_charts
                    )
                    mime_type = config.FORMATS_EXPORT["word"]["mime"]
                
                elif selected_format == "csv":
                    file_data, filename = generate_csv_export(
                        db_manager, export_type, export_config, user_session
                    )
                    mime_type = config.FORMATS_EXPORT["csv"]["mime"]
                
                elif selected_format == "pdf":
                    file_data, filename = generate_pdf_export(
                        db_manager, export_type, export_config, user_session, include_charts
                    )
                    mime_type = config.FORMATS_EXPORT["pdf"]["mime"]
                
                if file_data:
                    # Bouton de téléchargement
                    st.download_button(
                        label=f"📥 Télécharger {filename}",
                        data=file_data,
                        file_name=filename,
                        mime=mime_type,
                        type="secondary"
                    )
                    
                    FormComponents.alert_message(
                        f"Export {selected_format} généré avec succès !",
                        "success"
                    )
                else:
                    FormComponents.alert_message(
                        "Erreur lors de la génération de l'export",
                        "error"
                    )
        
        except Exception as e:
            logger.error(f"Erreur génération export: {e}")
            FormComponents.alert_message(f"Erreur export: {e}", "error")

def show_custom_reports(db_manager, db_utils, user_session: Dict):
    """Rapports personnalisés avec filtres avancés"""
    
    st.subheader("📝 Rapports Personnalisés")
    
    with st.form("custom_report_form"):
        st.markdown("#### ⚙️ Configuration du Rapport")
        
        # Sections à inclure
        col1, col2 = st.columns(2)
        
        with col1:
            include_operations = st.checkbox("📋 Vue opérations", value=True)
            include_timeline = st.checkbox("🎯 Timeline", value=False)
            include_risks = st.checkbox("🚨 Analyse risques", value=True)
            include_budget = st.checkbox("💰 Suivi budgétaire", value=True)
        
        with col2:
            include_rem = st.checkbox("📈 REM", value=False)
            include_alerts = st.checkbox("⚠️ Alertes", value=True)
            include_journal = st.checkbox("📖 Journal (extrait)", value=False)
            include_stats = st.checkbox("📊 Statistiques", value=True)
        
        # Filtres
        st.markdown("#### 🔍 Filtres")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Filtre par type d'opération
            selected_types = st.multiselect(
                "Types d'opération",
                options=config.TYPES_OPERATIONS,
                default=config.TYPES_OPERATIONS
            )
            
            # Filtre par statut
            selected_statuts = st.multiselect(
                "Statuts",
                options=list(config.STATUTS_OPERATIONS.keys()),
                default=list(config.STATUTS_OPERATIONS.keys()),
                format_func=lambda x: config.STATUTS_OPERATIONS[x]['libelle']
            )
        
        with col2:
            # Période
            date_range = st.date_input(
                "Période d'analyse",
                value=[date.today() - timedelta(days=90), date.today()],
                help="Période pour l'analyse des données"
            )
            
            # Score de risque minimum
            min_risk_score = st.slider(
                "Score de risque minimum",
                min_value=0,
                max_value=100,
                value=0,
                help="Inclure seulement les opérations avec score ≥"
            )
        
        with col3:
            # Format de rapport
            report_format = st.selectbox(
                "Format du rapport",
                options=["word", "pdf"],
                format_func=lambda x: "📄 Word" if x == "word" else "📑 PDF"
            )
            
            # Style de rapport
            report_style = st.selectbox(
                "Style de rapport",
                options=["executif", "detaille", "technique"],
                format_func=lambda x: {
                    "executif": "👔 Exécutif (synthèse)",
                    "detaille": "📋 Détaillé (complet)",
                    "technique": "🔧 Technique (analyse)"
                }[x]
            )
        
        # Métadonnées du rapport
        st.markdown("#### 📝 Informations du Rapport")
        
        col1, col2 = st.columns(2)
        
        with col1:
            report_title = st.text_input(
                "Titre du rapport",
                value=f"Rapport SPIC - {datetime.now().strftime('%d/%m/%Y')}"
            )
            
            report_author = st.text_input(
                "Auteur",
                value=user_session.get('nom', '') + " " + user_session.get('prenom', '')
            )
        
        with col2:
            report_description = st.text_area(
                "Description",
                value="Rapport automatique généré par SPIC 2.0",
                height=100
            )
        
        # Bouton génération
        submitted = st.form_submit_button("📋 Générer le Rapport Personnalisé", type="primary")
    
    if submitted:
        try:
            with st.spinner("Génération du rapport personnalisé en cours..."):
                
                # Configuration du rapport
                report_config = {
                    'sections': {
                        'operations': include_operations,
                        'timeline': include_timeline,
                        'risks': include_risks,
                        'budget': include_budget,
                        'rem': include_rem,
                        'alerts': include_alerts,
                        'journal': include_journal,
                        'stats': include_stats
                    },
                    'filters': {
                        'types': selected_types,
                        'statuts': selected_statuts,
                        'date_range': date_range,
                        'min_risk_score': min_risk_score
                    },
                    'metadata': {
                        'title': report_title,
                        'author': report_author,
                        'description': report_description,
                        'style': report_style
                    }
                }
                
                # Générer le rapport
                if report_format == "word":
                    file_data, filename = generate_custom_word_report(
                        db_manager, report_config, user_session
                    )
                    mime_type = config.FORMATS_EXPORT["word"]["mime"]
                else:
                    file_data, filename = generate_custom_pdf_report(
                        db_manager, report_config, user_session
                    )
                    mime_type = config.FORMATS_EXPORT["pdf"]["mime"]
                
                if file_data:
                    st.download_button(
                        label=f"📥 Télécharger {filename}",
                        data=file_data,
                        file_name=filename,
                        mime=mime_type,
                        type="secondary"
                    )
                    
                    FormComponents.alert_message(
                        "Rapport personnalisé généré avec succès !",
                        "success"
                    )
        
        except Exception as e:
            logger.error(f"Erreur génération rapport personnalisé: {e}")
            FormComponents.alert_message(f"Erreur génération rapport: {e}", "error")

def show_templates_management(db_manager, db_utils, user_session: Dict):
    """Gestion des templates d'export"""
    
    st.subheader("📋 Gestion des Templates")
    
    # Templates disponibles
    st.markdown("### 📄 Templates Disponibles")
    
    templates = get_available_templates()
    
    for template in templates:
        with st.expander(f"📝 {template['name']} - {template['description']}"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.write(f"**Type:** {template['type']}")
                st.write(f"**Format:** {template['format']}")
            
            with col2:
                st.write(f"**Sections:** {len(template['sections'])}")
                st.write(f"**Dernière MAJ:** {template['last_modified']}")
            
            with col3:
                if st.button(f"📥 Utiliser", key=f"use_template_{template['id']}"):
                    st.session_state['selected_template'] = template['id']
                    st.info(f"Template '{template['name']}' sélectionné")
                
                if st.button(f"👁️ Aperçu", key=f"preview_template_{template['id']}"):
                    show_template_preview(template)
    
    # Création de nouveau template
    st.markdown("### ➕ Créer un Nouveau Template")
    
    with st.expander("🛠️ Créateur de Template"):
        with st.form("new_template_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                template_name = st.text_input("Nom du template")
                template_type = st.selectbox(
                    "Type de template",
                    options=["operation", "portfolio", "risks", "financial"]
                )
            
            with col2:
                template_format = st.selectbox(
                    "Format",
                    options=["word", "excel"],
                    format_func=lambda x: "📄 Word" if x == "word" else "📊 Excel"
                )
                template_description = st.text_input("Description")
            
            # Sections à inclure
            st.markdown("**Sections à inclure:**")
            
            template_sections = {}
            col1, col2, col3 = st.columns(3)
            
            with col1:
                template_sections['header'] = st.checkbox("En-tête", value=True)
                template_sections['summary'] = st.checkbox("Résumé", value=True)
                template_sections['operations'] = st.checkbox("Opérations", value=True)
            
            with col2:
                template_sections['timeline'] = st.checkbox("Timeline", value=False)
                template_sections['budget'] = st.checkbox("Budget", value=True)
                template_sections['risks'] = st.checkbox("Risques", value=True)
            
            with col3:
                template_sections['alerts'] = st.checkbox("Alertes", value=True)
                template_sections['charts'] = st.checkbox("Graphiques", value=True)
                template_sections['footer'] = st.checkbox("Pied de page", value=True)
            
            submitted = st.form_submit_button("💾 Créer le Template")
            
            if submitted and template_name:
                # Sauvegarder le template (simulation)
                new_template = {
                    'name': template_name,
                    'type': template_type,
                    'format': template_format,
                    'description': template_description,
                    'sections': template_sections,
                    'created_by': user_session.get('user_id'),
                    'created_at': datetime.now().isoformat()
                }
                
                # Dans une vraie application, sauvegarder en BDD
                FormComponents.alert_message(
                    f"Template '{template_name}' créé avec succès !",
                    "success"
                )

def show_graphic_exports(db_manager, db_utils, user_session: Dict):
    """Exports spécifiques aux graphiques"""
    
    st.subheader("📈 Exports Graphiques")
    
    # Types de graphiques disponibles
    graphic_types = {
        "dashboard_metrics": "📊 Métriques Dashboard",
        "risk_distribution": "🚨 Distribution des Risques",
        "budget_evolution": "💰 Évolution Budgétaire",
        "timeline_gantt": "📅 Diagramme de Gantt",
        "rem_analysis": "📈 Analyse REM",
        "alerts_summary": "⚠️ Synthèse Alertes"
    }
    
    selected_graphic = st.selectbox(
        "Type de graphique",
        options=list(graphic_types.keys()),
        format_func=lambda x: graphic_types[x],
        key="graphic_export_type"
    )
    
    # Configuration selon le type de graphique
    graphic_config = configure_graphic_export(selected_graphic, db_utils, user_session)
    
    if graphic_config:
        # Formats d'export graphique
        col1, col2 = st.columns(2)
        
        with col1:
            export_format = st.selectbox(
                "Format d'export",
                options=["png", "pdf", "svg", "html"],
                format_func=lambda x: {
                    "png": "🖼️ PNG (image)",
                    "pdf": "📑 PDF",
                    "svg": "📐 SVG (vectoriel)",
                    "html": "🌐 HTML (interactif)"
                }[x]
            )
        
        with col2:
            if export_format in ["png", "pdf"]:
                resolution = st.selectbox(
                    "Résolution",
                    options=["standard", "high", "print"],
                    format_func=lambda x: {
                        "standard": "Standard (96 DPI)",
                        "high": "Haute (300 DPI)",
                        "print": "Impression (600 DPI)"
                    }[x]
                )
        
        # Aperçu du graphique
        if st.checkbox("👁️ Aperçu du graphique", value=True):
            try:
                preview_fig = generate_graphic_preview(
                    db_manager, selected_graphic, graphic_config, user_session
                )
                
                if preview_fig:
                    st.plotly_chart(preview_fig, use_container_width=True)
            
            except Exception as e:
                logger.error(f"Erreur aperçu graphique: {e}")
                st.error(f"Erreur génération aperçu: {e}")
        
        # Export du graphique
        if st.button("📥 Exporter le Graphique", type="primary"):
            try:
                with st.spinner("Export du graphique en cours..."):
                    
                    file_data, filename = export_graphic(
                        db_manager, selected_graphic, graphic_config, 
                        export_format, user_session
                    )
                    
                    if file_data:
                        mime_types = {
                            "png": "image/png",
                            "pdf": "application/pdf",
                            "svg": "image/svg+xml",
                            "html": "text/html"
                        }
                        
                        st.download_button(
                            label=f"📥 Télécharger {filename}",
                            data=file_data,
                            file_name=filename,
                            mime=mime_types.get(export_format, "application/octet-stream")
                        )
                        
                        FormComponents.alert_message(
                            f"Graphique exporté en {export_format.upper()} avec succès !",
                            "success"
                        )
            
            except Exception as e:
                logger.error(f"Erreur export graphique: {e}")
                FormComponents.alert_message(f"Erreur export graphique: {e}", "error")

# Fonctions utilitaires cohérentes avec l'architecture

def configure_export(export_type: str, db_utils: DatabaseUtils, user_session: Dict) -> Optional[Dict]:
    """Configurer les paramètres d'export selon le type"""
    
    config = {"type": export_type}
    
    if export_type == "operation_detail":
        # Sélection d'une opération spécifique
        operations_options = db_utils.get_operations_for_selectbox(user_session.get('user_id'))
        
        if not operations_options:
            st.warning("Aucune opération disponible")
            return None
        
        selected_op = st.selectbox(
            "Opération à exporter",
            options=list(operations_options.keys()),
            format_func=lambda x: operations_options[x],
            key="export_operation_select"
        )
        
        config["operation_id"] = selected_op
    
    elif export_type in ["journal", "risks", "alerts"]:
        # Filtres de période
        col1, col2 = st.columns(2)
        
        with col1:
            start_date = st.date_input(
                "Date de début",
                value=date.today() - timedelta(days=30),
                key=f"{export_type}_start_date"
            )
        
        with col2:
            end_date = st.date_input(
                "Date de fin",
                value=date.today(),
                key=f"{export_type}_end_date"
            )
        
        config["date_range"] = [start_date, end_date]
    
    elif export_type == "portfolio":
        # Filtres portefeuille
        col1, col2 = st.columns(2)
        
        with col1:
            selected_types = st.multiselect(
                "Types d'opération",
                options=config.TYPES_OPERATIONS,
                default=config.TYPES_OPERATIONS,
                key="portfolio_types"
            )
        
        with col2:
            selected_statuts = st.multiselect(
                "Statuts",
                options=list(config.STATUTS_OPERATIONS.keys()),
                default=list(config.STATUTS_OPERATIONS.keys()),
                format_func=lambda x: config.STATUTS_OPERATIONS[x]['libelle'],
                key="portfolio_statuts"
            )
        
        config["filters"] = {
            "types": selected_types,
            "statuts": selected_statuts
        }
    
    return config

def get_available_formats(export_type: str) -> List[str]:
    """Formats disponibles selon le type d'export"""
    
    format_mapping = {
        "portfolio": ["excel", "csv", "word"],
        "operation_detail": ["word", "pdf", "excel"],
        "journal": ["excel", "csv"],
        "risks": ["word", "pdf", "excel"],
        "budget_evolution": ["excel", "word"],
        "rem_summary": ["excel", "word"],
        "alerts": ["excel", "csv", "word"]
    }
    
    return format_mapping.get(export_type, ["excel", "csv"])

def show_export_preview(db_manager, export_type: str, export_config: Dict, user_session: Dict):
    """Aperçu des données avant export"""
    
    try:
        if export_type == "portfolio":
            # Récupérer données portefeuille
            operations = db_manager.get_all_operations(user_session.get('user_id'))
            if operations:
                df = pd.DataFrame(operations)
                
                # Appliquer filtres
                filters = export_config.get('filters', {})
                if filters.get('types'):
                    df = df[df['type_operation'].isin(filters['types'])]
                if filters.get('statuts'):
                    df = df[df['statut_global'].isin(filters['statuts'])]
                
                st.dataframe(df.head(10), use_container_width=True)
                st.info(f"📊 {len(df)} opération(s) à exporter")
        
        elif export_type == "operation_detail":
            # Détail opération
            operation_id = export_config.get('operation_id')
            if operation_id:
                operation = db_manager.get_operation(operation_id)
                phases = db_manager.get_phases_by_operation(operation_id)
                
                if operation:
                    st.write("**Opération:**", operation['nom'])
                    st.write("**Type:**", operation['type_operation'])
                    st.write("**Phases:**", len(phases))
        
        elif export_type == "journal":
            # Aperçu journal
            date_range = export_config.get('date_range')
            if date_range:
                # Simulation aperçu journal
                st.info(f"📖 Journal du {date_range[0]} au {date_range[1]}")
        
        else:
            st.info(f"📋 Aperçu pour {export_type}")
    
    except Exception as e:
        logger.error(f"Erreur aperçu export: {e}")
        st.error(f"Erreur aperçu: {e}")

# Fonctions de génération d'exports

def generate_excel_export(db_manager, export_type: str, export_config: Dict, user_session: Dict) -> tuple:
    """Générer export Excel"""
    
    try:
        # Créer un buffer en mémoire
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            
            if export_type == "portfolio":
                # Export portefeuille
                operations = db_manager.get_all_operations(user_session.get('user_id'))
                if operations:
                    df = pd.DataFrame(operations)
                    
                    # Appliquer filtres
                    filters = export_config.get('filters', {})
                    if filters.get('types'):
                        df = df[df['type_operation'].isin(filters['types'])]
                    if filters.get('statuts'):
                        df = df[df['statut_global'].isin(filters['statuts'])]
                    
                    # Formater colonnes
                    df['budget_initial'] = df['budget_initial'].fillna(0)
                    df['date_debut'] = pd.to_datetime(df['date_debut']).dt.strftime('%d/%m/%Y')
                    
                    df.to_excel(writer, sheet_name='Opérations', index=False)
            
            elif export_type == "operation_detail":
                # Export détail opération
                operation_id = export_config.get('operation_id')
                if operation_id:
                    # Feuille opération
                    operation = db_manager.get_operation(operation_id)
                    if operation:
                        op_df = pd.DataFrame([operation])
                        op_df.to_excel(writer, sheet_name='Opération', index=False)
                    
                    # Feuille phases
                    phases = db_manager.get_phases_by_operation(operation_id)
                    if phases:
                        phases_df = pd.DataFrame(phases)
                        phases_df.to_excel(writer, sheet_name='Phases', index=False)
                    
                    # Feuille budget
                    budgets = db_manager.get_budget_evolution(operation_id)
                    if budgets:
                        budget_df = pd.DataFrame(budgets)
                        budget_df.to_excel(writer, sheet_name='Budgets', index=False)
            
            elif export_type == "journal":
                # Export journal
                date_range = export_config.get('date_range')
                if date_range:
                    journal = db_manager.get_journal_entries(limit=1000)
                    if journal:
                        journal_df = pd.DataFrame(journal)
                        journal_df['timestamp'] = pd.to_datetime(journal_df['timestamp']).dt.strftime('%d/%m/%Y %H:%M')
                        journal_df.to_excel(writer, sheet_name='Journal', index=False)
        
        output.seek(0)
        filename = f"export_spic_{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        return output.getvalue(), filename
    
    except Exception as e:
        logger.error(f"Erreur génération Excel: {e}")
        return None, None

def generate_word_export(db_manager, export_type: str, export_config: Dict, 
                        user_session: Dict, include_charts: bool = False) -> tuple:
    """Générer export Word avec python-docx"""
    
    try:
        # Créer document Word
        doc = Document()
        
        # En-tête du document
        title = doc.add_heading(f'Rapport SPIC - {export_type.title()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations générales
        doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        doc.add_paragraph(f"Utilisateur : {user_session.get('nom', '')} {user_session.get('prenom', '')}")
        doc.add_paragraph("")
        
        if export_type == "portfolio":
            # Rapport portefeuille
            generate_portfolio_word_section(doc, db_manager, export_config, user_session)
        
        elif export_type == "operation_detail":
            # Rapport détail opération
            operation_id = export_config.get('operation_id')
            if operation_id:
                generate_operation_detail_word_section(doc, db_manager, operation_id, include_charts)
        
        elif export_type == "risks":
            # Rapport risques
            generate_risks_word_section(doc, db_manager, export_config, user_session)
        
            # Sauvegarder en mémoire
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        filename = f"rapport_spic_{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return output.getvalue(), filename
    
    except Exception as e:
        logger.error(f"Erreur génération Word: {e}")
        return None, None

def generate_csv_export(db_manager, export_type: str, export_config: Dict, user_session: Dict) -> tuple:
    """Générer export CSV"""
    
    try:
        if export_type == "portfolio":
            operations = db_manager.get_all_operations(user_session.get('user_id'))
            if operations:
                df = pd.DataFrame(operations)
                
                # Appliquer filtres
                filters = export_config.get('filters', {})
                if filters.get('types'):
                    df = df[df['type_operation'].isin(filters['types'])]
                if filters.get('statuts'):
                    df = df[df['statut_global'].isin(filters['statuts'])]
                
                csv_data = df.to_csv(index=False, encoding='utf-8-sig', sep=';')
        
        elif export_type == "journal":
            journal = db_manager.get_journal_entries(limit=1000)
            if journal:
                df = pd.DataFrame(journal)
                df['timestamp'] = pd.to_datetime(df['timestamp']).dt.strftime('%d/%m/%Y %H:%M')
                csv_data = df.to_csv(index=False, encoding='utf-8-sig', sep=';')
        
        elif export_type == "alerts":
            alerts = db_manager.get_active_alerts()
            if alerts:
                df = pd.DataFrame(alerts)
                df['date_creation'] = pd.to_datetime(df['date_creation']).dt.strftime('%d/%m/%Y')
                csv_data = df.to_csv(index=False, encoding='utf-8-sig', sep=';')
        
        else:
            csv_data = "Type d'export CSV non supporté"
        
        filename = f"export_spic_{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        
        return csv_data.encode('utf-8-sig'), filename
    
    except Exception as e:
        logger.error(f"Erreur génération CSV: {e}")
        return None, None

def generate_pdf_export(db_manager, export_type: str, export_config: Dict, 
                       user_session: Dict, include_charts: bool = False) -> tuple:
    """Générer export PDF (via conversion Word)"""
    
    try:
        # Pour simplifier, on génère d'abord un Word puis on le convertit
        # Dans une vraie application, on utiliserait reportlab ou weasyprint
        
        word_data, word_filename = generate_word_export(
            db_manager, export_type, export_config, user_session, include_charts
        )
        
        if word_data:
            # Simulation conversion PDF
            # Dans la réalité : utiliser python-docx2pdf ou similaire
            pdf_filename = word_filename.replace('.docx', '.pdf')
            
            # Pour l'instant, retourner le Word comme PDF (simulation)
            return word_data, pdf_filename
        
        return None, None
    
    except Exception as e:
        logger.error(f"Erreur génération PDF: {e}")
        return None, None

# Fonctions de génération des sections Word

def generate_portfolio_word_section(doc, db_manager, export_config: Dict, user_session: Dict):
    """Générer section portefeuille dans document Word"""
    
    try:
        doc.add_heading('Portefeuille d\'Opérations', level=1)
        
        # Récupérer données
        operations = db_manager.get_all_operations(user_session.get('user_id'))
        
        if not operations:
            doc.add_paragraph("Aucune opération trouvée.")
            return
        
        df = pd.DataFrame(operations)
        
        # Appliquer filtres
        filters = export_config.get('filters', {})
        if filters.get('types'):
            df = df[df['type_operation'].isin(filters['types'])]
        if filters.get('statuts'):
            df = df[df['statut_global'].isin(filters['statuts'])]
        
        # Statistiques générales
        doc.add_heading('Statistiques Générales', level=2)
        doc.add_paragraph(f"Nombre total d'opérations : {len(df)}")
        
        # Répartition par type
        type_counts = df['type_operation'].value_counts()
        doc.add_paragraph("Répartition par type d'opération :")
        for op_type, count in type_counts.items():
            doc.add_paragraph(f"  • {op_type} : {count} opération(s)", style='List Bullet')
        
        # Répartition par statut
        status_counts = df['statut_global'].value_counts()
        doc.add_paragraph("Répartition par statut :")
        for status, count in status_counts.items():
            status_label = config.STATUTS_OPERATIONS.get(status, {}).get('libelle', status)
            doc.add_paragraph(f"  • {status_label} : {count} opération(s)", style='List Bullet')
        
        # Budget total
        budget_total = df['budget_initial'].fillna(0).sum()
        doc.add_paragraph(f"Budget total du portefeuille : {format_number_french(budget_total)}€")
        
        # Score de risque moyen
        score_moyen = df['score_risque'].fillna(0).mean()
        doc.add_paragraph(f"Score de risque moyen : {score_moyen:.1f}/100")
        
        # Table des opérations
        doc.add_heading('Liste des Opérations', level=2)
        
        # Créer tableau
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # En-têtes
        headers = ['Nom', 'Type', 'Statut', 'Budget Initial', 'Score Risque']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        
        # Données
        for _, operation in df.iterrows():
            row = table.add_row()
            row.cells[0].text = operation.get('nom', '')
            row.cells[1].text = operation.get('type_operation', '')
            row.cells[2].text = config.STATUTS_OPERATIONS.get(
                operation.get('statut_global', ''), {}
            ).get('libelle', operation.get('statut_global', ''))
            row.cells[3].text = f"{format_number_french(operation.get('budget_initial', 0))}€"
            row.cells[4].text = f"{operation.get('score_risque', 0)}/100"
    
    except Exception as e:
        logger.error(f"Erreur section portefeuille Word: {e}")
        doc.add_paragraph(f"Erreur génération section portefeuille: {e}")

def generate_operation_detail_word_section(doc, db_manager, operation_id: str, include_charts: bool):
    """Générer section détail opération dans document Word"""
    
    try:
        operation = db_manager.get_operation(operation_id)
        
        if not operation:
            doc.add_paragraph("Opération non trouvée.")
            return
        
        doc.add_heading(f'Détail de l\'Opération: {operation["nom"]}', level=1)
        
        # Informations générales
        doc.add_heading('Informations Générales', level=2)
        
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'
        
        infos = [
            ('Nom', operation.get('nom', '')),
            ('Type d\'opération', operation.get('type_operation', '')),
            ('Statut', config.STATUTS_OPERATIONS.get(operation.get('statut_global', ''), {}).get('libelle', '')),
            ('Budget initial', f"{format_number_french(operation.get('budget_initial', 0))}€"),
            ('Surface', f"{format_number_french(operation.get('surface_m2', 0))} m²" if operation.get('surface_m2') else 'Non définie'),
            ('Secteur géographique', operation.get('secteur_geographique', 'Non défini')),
            ('Score de risque', f"{operation.get('score_risque', 0)}/100"),
            ('Date de création', operation.get('created_at', '')[:10] if operation.get('created_at') else '')
        ]
        
        for i, (label, value) in enumerate(infos):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            info_table.cell(i, 1).text = str(value)
        
        # Phases
        phases = db_manager.get_phases_by_operation(operation_id)
        
        if phases:
            doc.add_heading('Phases de l\'Opération', level=2)
            
            phases_table = doc.add_table(rows=1, cols=4)
            phases_table.style = 'Table Grid'
            
            # En-têtes phases
            phase_headers = ['Phase', 'Statut', 'Progression', 'Principale']
            for i, header in enumerate(phase_headers):
                phases_table.cell(0, i).text = header
                phases_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
            # Données phases
            for phase in phases:
                row = phases_table.add_row()
                row.cells[0].text = phase.get('nom_phase', '')
                row.cells[1].text = config.STATUTS_PHASES.get(
                    phase.get('statut_phase', ''), {}
                ).get('libelle', phase.get('statut_phase', ''))
                row.cells[2].text = f"{phase.get('progression_pct', 0)}%"
                row.cells[3].text = "Oui" if phase.get('principale') else "Non"
        
        # Budget
        budgets = db_manager.get_budget_evolution(operation_id)
        
        if budgets:
            doc.add_heading('Évolution Budgétaire', level=2)
            
            budget_table = doc.add_table(rows=1, cols=4)
            budget_table.style = 'Table Grid'
            
            budget_headers = ['Type', 'Montant', 'Date', 'Justification']
            for i, header in enumerate(budget_headers):
                budget_table.cell(0, i).text = header
                budget_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
            for budget in budgets:
                row = budget_table.add_row()
                row.cells[0].text = budget.get('type_budget', '')
                row.cells[1].text = f"{format_number_french(budget.get('montant', 0))}€"
                row.cells[2].text = budget.get('date_budget', '')[:10] if budget.get('date_budget') else ''
                row.cells[3].text = budget.get('justification', '')[:50] + '...' if len(budget.get('justification', '')) > 50 else budget.get('justification', '')
        
        # Alertes actives
        alerts = db_manager.get_active_alerts(operation_id)
        
        if alerts:
            doc.add_heading('Alertes Actives', level=2)
            
            for alert in alerts:
                doc.add_paragraph(
                    f"• {alert.get('titre', 'Alerte')} ({alert.get('niveau_severite', 'MOYEN')})",
                    style='List Bullet'
                )
                if alert.get('description'):
                    doc.add_paragraph(f"  Description: {alert['description']}")
    
    except Exception as e:
        logger.error(f"Erreur section détail opération Word: {e}")
        doc.add_paragraph(f"Erreur génération section détail: {e}")

def generate_risks_word_section(doc, db_manager, export_config: Dict, user_session: Dict):
    """Générer section analyse des risques dans document Word"""
    
    try:
        doc.add_heading('Analyse des Risques', level=1)
        
        # TOP des risques
        top_risks = db_manager.get_top_risks(10)
        
        if top_risks:
            doc.add_heading('TOP 10 des Opérations à Risque', level=2)
            
            risks_table = doc.add_table(rows=1, cols=4)
            risks_table.style = 'Table Grid'
            
            risk_headers = ['Rang', 'Opération', 'Score de Risque', 'Statut']
            for i, header in enumerate(risk_headers):
                risks_table.cell(0, i).text = header
                risks_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
            for i, risk_op in enumerate(top_risks, 1):
                row = risks_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = risk_op.get('nom', '')
                row.cells[2].text = f"{risk_op.get('score_risque', 0)}/100"
                row.cells[3].text = config.STATUTS_OPERATIONS.get(
                    risk_op.get('statut_global', ''), {}
                ).get('libelle', risk_op.get('statut_global', ''))
        
        # Statistiques globales des risques
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT 
                    AVG(score_risque) as score_moyen,
                    COUNT(CASE WHEN score_risque >= 75 THEN 1 END) as critiques,
                    COUNT(CASE WHEN score_risque >= 50 THEN 1 END) as eleves,
                    COUNT(*) as total
                FROM operations
            """)
            
            stats = dict(cursor.fetchone())
        
        if stats:
            doc.add_heading('Statistiques des Risques', level=2)
            doc.add_paragraph(f"Score de risque moyen du portefeuille : {stats.get('score_moyen', 0):.1f}/100")
            doc.add_paragraph(f"Opérations critiques (≥75) : {stats.get('critiques', 0)}")
            doc.add_paragraph(f"Opérations à risque élevé (≥50) : {stats.get('eleves', 0)}")
            doc.add_paragraph(f"Total opérations analysées : {stats.get('total', 0)}")
        
        # Alertes actives
        alerts = db_manager.get_active_alerts()
        
        if alerts:
            doc.add_heading('Alertes Actives', level=2)
            
            # Grouper par niveau
            alerts_by_level = {}
            for alert in alerts:
                level = alert.get('niveau_severite', 'MOYEN')
                if level not in alerts_by_level:
                    alerts_by_level[level] = []
                alerts_by_level[level].append(alert)
            
            for level in ['CRITIQUE', 'ELEVE', 'MOYEN', 'FAIBLE']:
                if level in alerts_by_level:
                    doc.add_paragraph(f"{level} ({len(alerts_by_level[level])} alertes):", style='Heading 3')
                    
                    for alert in alerts_by_level[level][:5]:  # Limiter à 5 par niveau
                        doc.add_paragraph(
                            f"• {alert.get('titre', 'Alerte')} - {alert.get('operation_nom', 'Opération')}",
                            style='List Bullet'
                        )
    
    except Exception as e:
        logger.error(f"Erreur section risques Word: {e}")
        doc.add_paragraph(f"Erreur génération section risques: {e}")

# Fonctions pour rapports personnalisés

def generate_custom_word_report(db_manager, report_config: Dict, user_session: Dict) -> tuple:
    """Générer rapport Word personnalisé"""
    
    try:
        doc = Document()
        
        # Métadonnées du rapport
        metadata = report_config.get('metadata', {})
        
        # Titre
        title = doc.add_heading(metadata.get('title', 'Rapport SPIC Personnalisé'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations du rapport
        doc.add_paragraph(f"Auteur : {metadata.get('author', '')}")
        doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        doc.add_paragraph(f"Style : {metadata.get('style', 'Standard')}")
        doc.add_paragraph("")
        
        if metadata.get('description'):
            doc.add_paragraph(metadata['description'])
            doc.add_paragraph("")
        
        # Sections selon configuration
        sections = report_config.get('sections', {})
        filters = report_config.get('filters', {})
        
        if sections.get('operations'):
            # Adapter export_config pour les filtres
            export_config = {'filters': filters}
            generate_portfolio_word_section(doc, db_manager, export_config, user_session)
        
        if sections.get('risks'):
            generate_risks_word_section(doc, db_manager, {}, user_session)
        
        if sections.get('budget'):
            generate_budget_word_section(doc, db_manager, filters, user_session)
        
        if sections.get('alerts'):
            generate_alerts_word_section(doc, db_manager, filters, user_session)
        
        if sections.get('stats'):
            generate_stats_word_section(doc, db_manager, filters, user_session)
        
        # Sauvegarder
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        filename = f"rapport_personnalise_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return output.getvalue(), filename
    
    except Exception as e:
        logger.error(f"Erreur rapport personnalisé Word: {e}")
        return None, None

def generate_budget_word_section(doc, db_manager, filters: Dict, user_session: Dict):
    """Section budget pour rapport personnalisé"""
    
    try:
        doc.add_heading('Analyse Budgétaire', level=1)
        
        # Récupérer données budgétaires globales
        with db_manager.get_connection() as conn:
            cursor = conn.cursor()
            
            where_conditions = ["1=1"]
            params = []
            
            # Appliquer filtres
            if filters.get('types'):
                placeholders = ','.join('?' for _ in filters['types'])
                where_conditions.append(f"type_operation IN ({placeholders})")
                params.extend(filters['types'])
            
            query = f"""
                SELECT 
                    SUM(budget_initial) as budget_initial_total,
                    SUM(COALESCE(budget_final, budget_revise, budget_initial)) as budget_actuel_total,
                    COUNT(*) as nb_operations,
                    AVG(budget_initial) as budget_moyen
                FROM operations
                WHERE {' AND '.join(where_conditions)}
            """
            
            cursor.execute(query, params)
            budget_stats = dict(cursor.fetchone())
        
        if budget_stats:
            doc.add_paragraph(f"Budget initial total : {format_number_french(budget_stats.get('budget_initial_total', 0))}€")
            doc.add_paragraph(f"Budget actuel total : {format_number_french(budget_stats.get('budget_actuel_total', 0))}€")
            
            if budget_stats.get('budget_initial_total', 0) > 0:
                variation = budget_stats['budget_actuel_total'] - budget_stats['budget_initial_total']
                variation_pct = (variation / budget_stats['budget_initial_total']) * 100
                doc.add_paragraph(f"Variation budgétaire : {variation_pct:+.1f}% ({format_number_french(variation):+}€)")
            
            doc.add_paragraph(f"Budget moyen par opération : {format_number_french(budget_stats.get('budget_moyen', 0))}€")
    
    except Exception as e:
        logger.error(f"Erreur section budget: {e}")
        doc.add_paragraph(f"Erreur génération section budget: {e}")

def generate_alerts_word_section(doc, db_manager, filters: Dict, user_session: Dict):
    """Section alertes pour rapport personnalisé"""
    
    try:
        doc.add_heading('Alertes et Surveillance', level=1)
        
        alerts = db_manager.get_active_alerts()
        
        if not alerts:
            doc.add_paragraph("Aucune alerte active.")
            return
        
        # Statistiques alertes
        alerts_by_level = {}
        for alert in alerts:
            level = alert.get('niveau_severite', 'MOYEN')
            if level not in alerts_by_level:
                alerts_by_level[level] = 0
            alerts_by_level[level] += 1
        
        doc.add_paragraph(f"Nombre total d'alertes actives : {len(alerts)}")
        
        for level in ['CRITIQUE', 'ELEVE', 'MOYEN', 'FAIBLE']:
            if level in alerts_by_level:
                doc.add_paragraph(f"  • {level} : {alerts_by_level[level]} alerte(s)", style='List Bullet')
        
        # Détail des alertes critiques
        critical_alerts = [a for a in alerts if a.get('niveau_severite') == 'CRITIQUE']
        
        if critical_alerts:
            doc.add_heading('Alertes Critiques', level=2)
            
            for alert in critical_alerts[:10]:  # Limiter à 10
                doc.add_paragraph(
                    f"• {alert.get('titre', 'Alerte')} - {alert.get('operation_nom', 'Opération')}",
                    style='List Bullet'
                )
                if alert.get('description'):
                    doc.add_paragraph(f"  {alert['description']}")
    
    except Exception as e:
        logger.error(f"Erreur section alertes: {e}")
        doc.add_paragraph(f"Erreur génération section alertes: {e}")

def generate_stats_word_section(doc, db_manager, filters: Dict, user_session: Dict):
    """Section statistiques pour rapport personnalisé"""
    
    try:
        doc.add_heading('Statistiques Générales', level=1)
        
        # Métriques dashboard
        dashboard_data = db_manager.get_operations_dashboard()
        
        if dashboard_data and dashboard_data.get('statistiques'):
            stats = dashboard_data['statistiques']
            
            doc.add_paragraph("Indicateurs clés du portefeuille :")
            doc.add_paragraph(f"  • Total opérations : {stats.get('total_operations', 0)}", style='List Bullet')
            doc.add_paragraph(f"  • Opérations actives : {stats.get('operations_actives', 0)}", style='List Bullet')
            doc.add_paragraph(f"  • Opérations terminées : {stats.get('operations_terminees', 0)}", style='List Bullet')
            doc.add_paragraph(f"  • Opérations bloquées : {stats.get('operations_bloquees', 0)}", style='List Bullet')
            doc.add_paragraph(f"  • Score de risque moyen : {stats.get('score_risque_moyen', 0):.1f}/100", style='List Bullet')
        
        # Répartition par type
        if dashboard_data and dashboard_data.get('repartition_type'):
            doc.add_paragraph("")
            doc.add_paragraph("Répartition par type d'opération :")
            
            for op_type, count in dashboard_data['repartition_type'].items():
                doc.add_paragraph(f"  • {op_type} : {count} opération(s)", style='List Bullet')
    
    except Exception as e:
        logger.error(f"Erreur section statistiques: {e}")
        doc.add_paragraph(f"Erreur génération section statistiques: {e}")

# Fonctions pour exports graphiques

def configure_graphic_export(graphic_type: str, db_utils: DatabaseUtils, user_session: Dict) -> Dict:
    """Configuration pour export graphique"""
    
    config = {"type": graphic_type}
    
    if graphic_type == "timeline_gantt":
        # Sélection opération pour Gantt
        operations_options = db_utils.get_operations_for_selectbox(user_session.get('user_id'))
        
        if operations_options:
            selected_op = st.selectbox(
                "Opération pour diagramme de Gantt",
                options=list(operations_options.keys()),
                format_func=lambda x: operations_options[x],
                key="gantt_operation_select"
            )
            config["operation_id"] = selected_op
    
    elif graphic_type in ["budget_evolution", "rem_analysis"]:
        # Période pour évolution
        col1, col2 = st.columns(2)
        
        with col1:
            start_date = st.date_input(
                "Date de début",
                value=date.today() - timedelta(days=365),
                key=f"{graphic_type}_start"
            )
        
        with col2:
            end_date = st.date_input(
                "Date de fin",
                value=date.today(),
                key=f"{graphic_type}_end"
            )
        
        config["date_range"] = [start_date, end_date]
    
    return config

def generate_graphic_preview(db_manager, graphic_type: str, graphic_config: Dict, user_session: Dict):
    """Générer aperçu du graphique"""
    
    try:
        if graphic_type == "dashboard_metrics":
            dashboard_data = db_manager.get_operations_dashboard()
            if dashboard_data:
                return create_dashboard_metrics(dashboard_data)
        
        elif graphic_type == "risk_distribution":
            from modules.gestion_risques import get_risk_distribution, create_risk_distribution_chart
            risk_data = get_risk_distribution(db_manager, user_session.get('user_id'))
            if risk_data:
                return create_risk_distribution_chart(risk_data)
        
        elif graphic_type == "budget_evolution":
            # Évolution budgétaire globale
            budget_evolution = db_manager.get_operations_dashboard()
            if budget_evolution:
                return create_budget_evolution_chart([])  # Données simplifiées
        
        elif graphic_type == "timeline_gantt":
            operation_id = graphic_config.get('operation_id')
            if operation_id:
                phases = db_manager.get_phases_by_operation(operation_id)
                if phases:
                    return create_gantt_chart(phases)
        
        elif graphic_type == "alerts_summary":
            from modules.gestion_risques import get_alerts_by_type, create_alerts_by_type_chart
            alerts_data = get_alerts_by_type(db_manager, user_session.get('user_id'))
            if alerts_data:
                return create_alerts_by_type_chart(alerts_data)
        
        return None
    
    except Exception as e:
        logger.error(f"Erreur aperçu graphique {graphic_type}: {e}")
        return None

def export_graphic(db_manager, graphic_type: str, graphic_config: Dict, 
                  export_format: str, user_session: Dict) -> tuple:
    """Exporter graphique dans le format demandé"""
    
    try:
        # Générer le graphique
        fig = generate_graphic_preview(db_manager, graphic_type, graphic_config, user_session)
        
        if not fig:
            return None, None
        
        # Export selon format
        if export_format == "png":
            img_bytes = fig.to_image(format="png", width=1200, height=800, scale=2)
            filename = f"graphique_{graphic_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            return img_bytes, filename
        
        elif export_format == "pdf":
            img_bytes = fig.to_image(format="pdf", width=1200, height=800)
            filename = f"graphique_{graphic_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            return img_bytes, filename
        
        elif export_format == "svg":
            svg_bytes = fig.to_image(format="svg", width=1200, height=800)
            filename = f"graphique_{graphic_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.svg"
            return svg_bytes, filename
        
        elif export_format == "html":
            html_str = fig.to_html(include_plotlyjs='cdn')
            filename = f"graphique_{graphic_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            return html_str.encode('utf-8'), filename
        
        return None, None
    
    except Exception as e:
        logger.error(f"Erreur export graphique: {e}")
        return None, None

# Templates

def get_available_templates() -> List[Dict]:
    """Liste des templates disponibles"""
    
    return [
        {
            'id': 'template_1',
            'name': 'Rapport Opération Standard',
            'type': 'operation',
            'format': 'word',
            'description': 'Rapport détaillé pour une opération',
            'sections': ['header', 'summary', 'phases', 'budget', 'risks'],
            'last_modified': '15/12/2024'
        },
        {
            'id': 'template_2', 
            'name': 'Synthèse Portefeuille',
            'type': 'portfolio',
            'format': 'excel',
            'description': 'Vue d\'ensemble du portefeuille',
            'sections': ['operations', 'stats', 'charts'],
            'last_modified': '10/12/2024'
        },
        {
            'id': 'template_3',
            'name': 'Analyse Risques Exécutif',
            'type': 'risks',
            'format': 'word',
            'description': 'Rapport risques pour direction',
            'sections': ['summary', 'top_risks', 'recommendations'],
            'last_modified': '08/12/2024'
        },
        {
            'id': 'template_4',
            'name': 'Tableau de Bord Financier',
            'type': 'financial',
            'format': 'excel',
            'description': 'Suivi budgétaire et REM',
            'sections': ['budget_evolution', 'rem_summary', 'alerts'],
            'last_modified': '05/12/2024'
        }
    ]

def show_template_preview(template: Dict):
    """Afficher aperçu d'un template"""
    
    st.markdown(f"### 👁️ Aperçu - {template['name']}")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.write(f"**Type:** {template['type']}")
        st.write(f"**Format:** {template['format'].upper()}")
        st.write(f"**Description:** {template['description']}")
    
    with col2:
        st.write("**Sections incluses:**")
        for section in template['sections']:
            section_names = {
                'header': '📄 En-tête',
                'summary': '📋 Résumé',
                'operations': '🏗️ Opérations',
                'phases': '🎯 Phases',
                'budget': '💰 Budget',
                'budget_evolution': '📈 Évolution Budget',
                'risks': '🚨 Risques',
                'top_risks': '🏆 TOP Risques',
                'stats': '📊 Statistiques',
                'charts': '📈 Graphiques',
                'alerts': '⚠️ Alertes',
                'rem_summary': '📈 Synthèse REM',
                'recommendations': '💡 Recommandations'
            }
            
            section_name = section_names.get(section, section)
            st.write(f"• {section_name}")
    
    # Exemple de contenu
    if template['type'] == 'operation':
        st.markdown("**Exemple de contenu:**")
        st.code("""
1. Informations Générales
   - Nom de l'opération
   - Type et statut
   - Budget et surface
   
2. Phases
   - Liste des phases avec statuts
   - Progression et retards
   
3. Analyse des Risques
   - Score de risque
   - Alertes actives
   
4. Suivi Budgétaire
   - Évolution du budget
   - Écarts et justifications
        """, language=None)

def generate_custom_pdf_report(db_manager, report_config: Dict, user_session: Dict) -> tuple:
    """Générer rapport PDF personnalisé (via conversion Word)"""
    
    try:
        # Pour simplifier, utiliser la génération Word puis conversion
        word_data, word_filename = generate_custom_word_report(
            db_manager, report_config, user_session
        )
        
        if word_data:
            pdf_filename = word_filename.replace('.docx', '.pdf')
            
            # Dans une vraie application, conversion Word vers PDF
            # Utiliser python-docx2pdf, pandoc, ou LibreOffice en ligne de commande
            
            # Pour l'instant, retourner le Word comme PDF (simulation)
            return word_data, pdf_filename
        
        return None, None
    
    except Exception as e:
        logger.error(f"Erreur génération PDF personnalisé: {e}")
        return None, None